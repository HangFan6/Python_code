# -*- coding:utf-8 -*-
"""
作者：HET
    简历筛选
日期：2023年04月12日
"""
from docx import Document
import glob


class ReadDoc(object):
    def __init__(self, path):
        self.doc = Document(path)
        self.p_text = ''
        self.table_text = ''
        self.textbox=''

        self.get_para()
        self.get_table()
        self.get_textbox()

    # 段落读取
    def get_para(self):
        for p in self.doc.paragraphs:
            self.p_text += p.text + '\n'  # 注意：有的Word文档看似是段落文档，其实是表格文档

    # 表格读取
    def get_table(self):
        for table in self.doc.tables:
            for row in table.rows:
                _cell_str = ''
                for cell in row.cells:
                    _cell_str += cell.text + ','  # 显示每行内容
                self.table_text += _cell_str + '\n'

    # 文本框读取
    # ======（待完善：多个文本框组合的内容读取）==================
    def get_textbox(self):
        children = self.doc.element.body.iter()
        count = 0  # 用于定位是哪个文本框
        for child in children:
            # 通过类型判断目录
            if child.tag.endswith('txbx'):
                for ci in child.iter():
                    if ci.tag.endswith('main}r'):
                        count += 1
                        self.textbox=ci.text

def search_word(path, targets):
    result = glob.glob(path)
    finall_result=[]
    for i in result:
        # 判断是否是文件
        isuse=True
        if glob.os.path.isfile(i):
            # 判断是否为docx文件
            if i.endswith('.docx'):
                # print(i)
                doc = ReadDoc(i)
                p_text = doc.p_text  # 读取文档段落
                t_text = doc.table_text  # 读取文档表格
                textbox = doc.textbox  # 读取文档表格
                all_text = p_text + t_text+textbox
                # print(textbox)
                # all_text = p_text + t_text
                for target in targets:
                    if target not in all_text:
                        isuse=False
                        break
                if not isuse:
                    continue
                finall_result.append(i)
    return finall_result



if __name__ == '__main__':
    path=glob.os.path.join(glob.os.getcwd(),'*')
    targets=['Python']   # 筛选关键字
    res=search_word(path,targets)
    print(res)
