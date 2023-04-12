# -*- coding:utf-8 -*-
"""
作者：HET
日期：2023年04月11日
"""
from docx import Document

# =======段落内容读取==========
doc=Document('Word.docx')  # doc格式需要转换成docx
# print(doc.paragraphs)  # 读取段落  返回列表
for p in doc.paragraphs:
    print(p.text)  # 读取段落中的信息

# =======表格内容读取==========
# 获取文档中每个表格对象t
for t in doc.tables:
    # 通过行的形式读取每个表格  列的形式使用columns函数
    for row in t.rows:
        _row_str=''
        # 读取行中的每个小表格内容
        for cell in row.cells:
            _row_str += cell.text + ','
            # print(cell.text)
        print(_row_str)  # 打印每行表格中的信息
    # print(dir(t))  # 查看表格中的函数





