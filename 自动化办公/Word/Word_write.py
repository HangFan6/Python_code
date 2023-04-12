# -*- coding:utf-8 -*-
"""
作者：HET
日期：2023年04月12日
"""
from docx import Document
from docx.shared import Inches,RGBColor,Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

doc=Document()
# ==========全局样式定义============
style=doc.styles['Normal']  # 设置文本格式
style.font.name='微软雅黑'  # 字体样式
"""
由于涉及操作系统和office版本等因素，如果中文字体不生效，
要导入模块from docx.oxml.ns import qn，
需借助style.element.rPr.rFonts.set(qn('w:eastAsia'), u'字体名')使之生效
"""
style.element.rPr.rFonts.set(qn('w:eastAsia'), u'华文楷体')

# style.font.color.rgb=RGBColor(255,125,20)  # 字体颜色
style.font.size=Pt(16)  # 字号大小

title=doc.add_heading('My title',0) # 创建标题，并使用样式0
title.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER  # 标题居中
title.style.font.size=Pt(20)
_t=title.add_run('\n123')    # 追加标题 并换行
_t.italic=True  # 字体倾斜  注意：只能将追加字体倾斜
_t.bold=True  # 字体加粗
# print(dir(_t))  # 查看样式函数  注意：只能将追加字体加粗

p=doc.add_paragraph('欢迎来到这里学习！')    # 添加段落内容
p.add_run('\n这是word生成的知识。').italic=True  # 段落内容追加
p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER  # 段落居中

# 添加图片
image_obj=doc.add_picture('../test3/images/imooc_1.jpg',width=Inches(5))   # 高可以通过宽自动化处理
# ======添加表格=========
title=['name','age','sex']
table=doc.add_table(rows=1,cols=3)
title_cells=table.rows[0].cells
title_cells[0].text=title[0]
title_cells[1].text=title[1]
title_cells[2].text=title[2]
# table中的内容必须是字符串
data=[
    ('xiaomu','10','man'),
    ('dewei','34','man'),
    ('xiaoman','18','women')
]
for d in data:
    # 追加到一行中的表格对象
    row_cells=table.add_row().cells
    row_cells[0].text=d[0]  # name
    row_cells[1].text=d[1]  # age
    row_cells[2].text=d[2]  # sex

# =====分页==========
doc.add_page_break()
title=doc.add_heading('My title 2',0)

doc.save('Create.docx')  # 保存Word文件
